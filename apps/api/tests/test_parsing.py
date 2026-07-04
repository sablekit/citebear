"""Parser tests: PDF font-size headings + page attribution, DOCX styles, dispatch.

Fixtures are generated in-process (fpdf2 for PDF, python-docx for DOCX) so the
suite has no binary blobs and the font sizes / styles under test are explicit.
"""

import io

import pytest
from docx import Document as DocxDocument
from fpdf import FPDF

from citebear_api.parsing import (
    PageLimitError,
    UnsupportedMediaTypeError,
    parse_document,
    parse_docx,
    parse_pdf,
)


def _pdf_with_heading_spanning_two_pages() -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=24)
    pdf.cell(0, 12, "Installation", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    for i in range(80):  # long enough to overflow onto a second page
        pdf.cell(
            0,
            6,
            f"Body line {i} about installing and configuring the package.",
            new_x="LMARGIN",
            new_y="NEXT",
        )
    return bytes(pdf.output())


def _docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Installation", level=1)
    document.add_paragraph("Run the installer to begin.")
    document.add_heading("Linux", level=2)
    document.add_paragraph("Use apt on Debian systems.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_pdf_detects_heading_from_font_size() -> None:
    sections, page_count = parse_pdf(_pdf_with_heading_spanning_two_pages())
    assert page_count >= 2
    installation = [s for s in sections if s.section_path == ["Installation"]]
    assert installation, "the larger-font line should become a heading"
    assert "# Installation" in installation[0].heading_line


def test_pdf_attributes_pages_across_a_boundary() -> None:
    sections, _ = parse_pdf(_pdf_with_heading_spanning_two_pages())
    section = next(s for s in sections if s.section_path == ["Installation"])
    # the body flows across the page break: it starts on page 1 and ends on 2
    assert section.page_start == 1
    assert section.page_end is not None and section.page_end >= 2


def test_pdf_body_is_not_all_headings() -> None:
    # a uniform-size body must not be misclassified as headings
    sections, _ = parse_pdf(_pdf_with_heading_spanning_two_pages())
    section = next(s for s in sections if s.section_path == ["Installation"])
    assert "Body line 0" in section.body


def _pdf_with_pages(count: int) -> bytes:
    pdf = FPDF()
    for page in range(count):
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)
        pdf.cell(0, 6, f"Page {page} body text.", new_x="LMARGIN", new_y="NEXT")
    return bytes(pdf.output())


def test_pdf_over_the_page_cap_is_rejected_mid_parse() -> None:
    # the cap must fire the moment the document crosses it, before the rest of
    # the pages are laid out (the whole point of enforcing it during parsing)
    with pytest.raises(PageLimitError):
        parse_pdf(_pdf_with_pages(3), max_pages=2)


def test_pdf_at_the_page_cap_is_accepted() -> None:
    sections, page_count = parse_pdf(_pdf_with_pages(2), max_pages=2)
    assert page_count == 2
    assert sections


def test_docx_reads_heading_styles_into_the_trail() -> None:
    sections, page_count = parse_docx(_docx_bytes())
    assert page_count is None  # DOCX stores no page numbers
    paths = [s.section_path for s in sections]
    assert ["Installation"] in paths
    assert ["Installation", "Linux"] in paths  # h2 nests under the preceding h1


def test_docx_sections_carry_no_pages() -> None:
    sections, _ = parse_docx(_docx_bytes())
    assert all(s.page_start is None and s.page_end is None for s in sections)


def test_parse_document_dispatches_markdown() -> None:
    sections, page_count = parse_document(b"# Title\n\nBody text.", "text/markdown")
    assert page_count is None
    assert sections[0].section_path == ["Title"]


def test_parse_document_rejects_unknown_mime() -> None:
    with pytest.raises(UnsupportedMediaTypeError):
        parse_document(b"\x89PNG...", "image/png")
